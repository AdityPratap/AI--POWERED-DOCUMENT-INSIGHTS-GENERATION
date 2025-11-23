<<<<<<< HEAD
# app.py — Streamlit Analyzer for PDFs & DOCX
import os
os.environ.setdefault("STREAMLIT_SERVER_FILE_WATCHER_TYPE", "none")

import io
import re
from typing import List, Tuple

import numpy as np
import streamlit as st
import PyPDF2
import docx  # for DOCX extraction

# Optional FAISS
try:
    import faiss  # type: ignore
    HAVE_FAISS = True
except Exception:
    HAVE_FAISS = False

from sentence_transformers import SentenceTransformer
from transformers import pipeline

# =============================
# App Config & Styles
# =============================
st.set_page_config(page_title="PDF/DOCX Analyzer", page_icon="📄", layout="wide")

st.markdown(
    """
    <style>
    html, body, [class*="css"] { font-size: 15px; }
    h1 { font-size: 1.6rem; }
    h2 { font-size: 1.3rem; }
    .muted { color: #6b7280; }
    .card { padding: 1rem; border: 1px solid #eee; border-radius: 14px; box-shadow: 0 1px 3px rgba(0,0,0,.06); background: #fff; }
    </style>
    """,
    unsafe_allow_html=True,
)

# =============================
# Models (cached)
# =============================
@st.cache_resource(show_spinner=False)
def get_embedder() -> SentenceTransformer:
    return SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")

@st.cache_resource(show_spinner=False)
def get_summarizer():
    return pipeline("summarization", model="facebook/bart-large-cnn")

# =============================
# Utilities
# =============================
def extract_text_from_pdf_bytes(file_bytes: bytes) -> str:
    """Extract text from PDF bytes."""
    text_parts: List[str] = []
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    for page in reader.pages:
        page_text = page.extract_text() or ""
        page_text = re.sub(r"\s+", " ", page_text).strip()
        if page_text:
            text_parts.append(page_text)
    return "\n\n".join(text_parts)

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes."""
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)

def embed_texts(texts: List[str]) -> np.ndarray:
    embedder = get_embedder()
    vectors = embedder.encode(texts, normalize_embeddings=True)
    return vectors.astype(np.float32)

def summarize_text(text: str, ratio: float = 0.1, max_tokens: int = 512) -> str:
    """Summarize text to a ratio of its length."""
    summarizer = get_summarizer()
    # Approx tokens from chars
    chunk_size = 1000
    chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
    summary_parts = []
    for ch in chunks:
        out = summarizer(ch, max_length=int(max_tokens*ratio), min_length=30, do_sample=False)
        summary_parts.append(out[0]['summary_text'])
    return " ".join(summary_parts)

# =============================
# Vector store (FAISS optional)
# =============================
class VectorStore:
    def __init__(self, dim: int):
        self.dim = dim
        self.ids: List[str] = []
        self.vectors: List[np.ndarray] = []
        if HAVE_FAISS:
            self.index = faiss.IndexFlatL2(dim)  # type: ignore
        else:
            self.index = None

    def add(self, vectors: np.ndarray, ids: List[str]):
        if vectors.ndim == 1:
            vectors = vectors.reshape(1, -1)
        if HAVE_FAISS and self.index is not None:
            self.index.add(vectors)
        self.vectors.extend([v for v in vectors])
        self.ids.extend(ids)

    def search(self, query_vec: np.ndarray, k: int = 5) -> List[Tuple[str, float]]:
        if query_vec.ndim == 1:
            query_vec = query_vec.reshape(1, -1)
        if HAVE_FAISS and self.index is not None and len(self.ids) > 0:
            D, I = self.index.search(query_vec, min(k, len(self.ids)))
            out: List[Tuple[str, float]] = []
            for idx, dist in zip(I.ravel().tolist(), D.ravel().tolist()):
                if 0 <= idx < len(self.ids):
                    out.append((self.ids[idx], float(dist)))
            return out
        if not self.vectors:
            return []
        M = np.vstack(self.vectors)
        q = query_vec[0]
        sims = M @ q
        dists = 1.0 - sims
        order = np.argsort(dists)[:k]
        return [(self.ids[i], float(dists[i])) for i in order]

# =============================
# Session state
# =============================
if "docs" not in st.session_state:
    st.session_state.docs = {}  # {doc_name: {"text": str, "summary10": str, "summary100": str}}
if "store" not in st.session_state:
    dim = get_embedder().get_sentence_embedding_dimension()
    st.session_state.store = VectorStore(dim)

# =============================
# UI
# =============================
st.title("📄 PDF & DOCX Analyzer with Summarization")

uploads = st.file_uploader("Upload files", type=["pdf", "docx"], accept_multiple_files=True)

if uploads:
    with st.spinner("Processing uploads…"):
        names, texts = [], []
        for uf in uploads:
            try:
                content = uf.read()
                if uf.name.lower().endswith(".pdf"):
                    text = extract_text_from_pdf_bytes(content)
                elif uf.name.lower().endswith(".docx"):
                    text = extract_text_from_docx(content)
                else:
                    st.warning(f"Unsupported file: {uf.name}")
                    continue

                if not text:
                    st.warning(f"No text detected in {uf.name}")
                    continue

                # Generate summaries
                short_sum = summarize_text(text, ratio=0.1)
                long_sum = summarize_text(text, ratio=1.0)

                st.session_state.docs[uf.name] = {
                    "text": text,
                    "summary10": short_sum,
                    "summary100": long_sum,
                }
                names.append(uf.name)
                texts.append(text)
            except Exception as e:
                st.error(f"Failed to process {uf.name}: {e}")
        if texts:
            vecs = embed_texts(texts)
            st.session_state.store.add(vecs, names)
    st.success("Upload and summarization complete.")

st.divider()

st.subheader("🔎 Query Summaries by Keyword")
user_q = st.text_input("Enter keywords or a question")

if user_q:
    qvec = embed_texts([user_q])
    results = st.session_state.store.search(qvec, k=min(3, len(st.session_state.store.ids)))
    if not results:
        st.info("No documents indexed yet.")
    else:
        st.write("**Top matches:**")
        for doc_id, dist in results:
            st.write(f"• {doc_id} — distance: {dist:.4f}")
        top_doc = results[0][0]
        st.subheader(f"📘 Best Match: {top_doc}")

        doc_data = st.session_state.docs.get(top_doc, {})
        short_sum = doc_data.get("summary10", "")
        long_sum = doc_data.get("summary100", "")

        # Filter summaries by keyword
        keyword = user_q.lower()
        filtered_short = " ".join([s for s in short_sum.split(". ") if keyword in s.lower()])
        filtered_long = " ".join([s for s in long_sum.split(". ") if keyword in s.lower()])

        st.markdown("**Short Summary (10%) filtered by keyword:**")
        st.text_area("Summary 10%", value=filtered_short or short_sum, height=200)

        st.markdown("**Detailed Summary (100%) filtered by keyword:**")
        st.text_area("Summary 100%", value=filtered_long or long_sum, height=300)

st.divider()

if st.session_state.docs:
    st.subheader("📚 Uploaded Documents")
    tabs = st.tabs(list(st.session_state.docs.keys()))
    for tab, name in zip(tabs, st.session_state.docs.keys()):
        with tab:
            st.text_area("Full Extracted Text", value=st.session_state.docs[name]["text"], height=300)
            st.download_button("⬇️ Download Extracted Text", data=st.session_state.docs[name]["text"], file_name=f"{name}_extracted.txt", mime="text/plain")
=======
# app.py — Streamlit Analyzer for PDFs & DOCX
import os
os.environ.setdefault("STREAMLIT_SERVER_FILE_WATCHER_TYPE", "none")

import io
import re
from typing import List, Tuple

import numpy as np
import streamlit as st
import PyPDF2
import docx  # for DOCX extraction

# Optional FAISS
try:
    import faiss  # type: ignore
    HAVE_FAISS = True
except Exception:
    HAVE_FAISS = False

from sentence_transformers import SentenceTransformer
from transformers import pipeline

# =============================
# App Config & Styles
# =============================
st.set_page_config(page_title="PDF/DOCX Analyzer", page_icon="📄", layout="wide")

st.markdown(
    """
    <style>
    html, body, [class*="css"] { font-size: 15px; }
    h1 { font-size: 1.6rem; }
    h2 { font-size: 1.3rem; }
    .muted { color: #6b7280; }
    .card { padding: 1rem; border: 1px solid #eee; border-radius: 14px; box-shadow: 0 1px 3px rgba(0,0,0,.06); background: #fff; }
    </style>
    """,
    unsafe_allow_html=True,
)

# =============================
# Models (cached)
# =============================
@st.cache_resource(show_spinner=False)
def get_embedder() -> SentenceTransformer:
    return SentenceTransformer("sentence-transformers/all-MiniLM-L6-v2")

@st.cache_resource(show_spinner=False)
def get_summarizer():
    return pipeline("summarization", model="facebook/bart-large-cnn")

# =============================
# Utilities
# =============================
def extract_text_from_pdf_bytes(file_bytes: bytes) -> str:
    """Extract text from PDF bytes."""
    text_parts: List[str] = []
    reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    for page in reader.pages:
        page_text = page.extract_text() or ""
        page_text = re.sub(r"\s+", " ", page_text).strip()
        if page_text:
            text_parts.append(page_text)
    return "\n\n".join(text_parts)

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes."""
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)

def embed_texts(texts: List[str]) -> np.ndarray:
    embedder = get_embedder()
    vectors = embedder.encode(texts, normalize_embeddings=True)
    return vectors.astype(np.float32)

def summarize_text(text: str, ratio: float = 0.1, max_tokens: int = 512) -> str:
    """Summarize text to a ratio of its length."""
    summarizer = get_summarizer()
    # Approx tokens from chars
    chunk_size = 1000
    chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
    summary_parts = []
    for ch in chunks:
        out = summarizer(ch, max_length=int(max_tokens*ratio), min_length=30, do_sample=False)
        summary_parts.append(out[0]['summary_text'])
    return " ".join(summary_parts)

# =============================
# Vector store (FAISS optional)
# =============================
class VectorStore:
    def __init__(self, dim: int):
        self.dim = dim
        self.ids: List[str] = []
        self.vectors: List[np.ndarray] = []
        if HAVE_FAISS:
            self.index = faiss.IndexFlatL2(dim)  # type: ignore
        else:
            self.index = None

    def add(self, vectors: np.ndarray, ids: List[str]):
        if vectors.ndim == 1:
            vectors = vectors.reshape(1, -1)
        if HAVE_FAISS and self.index is not None:
            self.index.add(vectors)
        self.vectors.extend([v for v in vectors])
        self.ids.extend(ids)

    def search(self, query_vec: np.ndarray, k: int = 5) -> List[Tuple[str, float]]:
        if query_vec.ndim == 1:
            query_vec = query_vec.reshape(1, -1)
        if HAVE_FAISS and self.index is not None and len(self.ids) > 0:
            D, I = self.index.search(query_vec, min(k, len(self.ids)))
            out: List[Tuple[str, float]] = []
            for idx, dist in zip(I.ravel().tolist(), D.ravel().tolist()):
                if 0 <= idx < len(self.ids):
                    out.append((self.ids[idx], float(dist)))
            return out
        if not self.vectors:
            return []
        M = np.vstack(self.vectors)
        q = query_vec[0]
        sims = M @ q
        dists = 1.0 - sims
        order = np.argsort(dists)[:k]
        return [(self.ids[i], float(dists[i])) for i in order]

# =============================
# Session state
# =============================
if "docs" not in st.session_state:
    st.session_state.docs = {}  # {doc_name: {"text": str, "summary10": str, "summary100": str}}
if "store" not in st.session_state:
    dim = get_embedder().get_sentence_embedding_dimension()
    st.session_state.store = VectorStore(dim)

# =============================
# UI
# =============================
st.title("📄 PDF & DOCX Analyzer with Summarization")

uploads = st.file_uploader("Upload files", type=["pdf", "docx"], accept_multiple_files=True)

if uploads:
    with st.spinner("Processing uploads…"):
        names, texts = [], []
        for uf in uploads:
            try:
                content = uf.read()
                if uf.name.lower().endswith(".pdf"):
                    text = extract_text_from_pdf_bytes(content)
                elif uf.name.lower().endswith(".docx"):
                    text = extract_text_from_docx(content)
                else:
                    st.warning(f"Unsupported file: {uf.name}")
                    continue

                if not text:
                    st.warning(f"No text detected in {uf.name}")
                    continue

                # Generate summaries
                short_sum = summarize_text(text, ratio=0.1)
                long_sum = summarize_text(text, ratio=1.0)

                st.session_state.docs[uf.name] = {
                    "text": text,
                    "summary10": short_sum,
                    "summary100": long_sum,
                }
                names.append(uf.name)
                texts.append(text)
            except Exception as e:
                st.error(f"Failed to process {uf.name}: {e}")
        if texts:
            vecs = embed_texts(texts)
            st.session_state.store.add(vecs, names)
    st.success("Upload and summarization complete.")

st.divider()

st.subheader("🔎 Query Summaries by Keyword")
user_q = st.text_input("Enter keywords or a question")

if user_q:
    qvec = embed_texts([user_q])
    results = st.session_state.store.search(qvec, k=min(3, len(st.session_state.store.ids)))
    if not results:
        st.info("No documents indexed yet.")
    else:
        st.write("**Top matches:**")
        for doc_id, dist in results:
            st.write(f"• {doc_id} — distance: {dist:.4f}")
        top_doc = results[0][0]
        st.subheader(f"📘 Best Match: {top_doc}")

        doc_data = st.session_state.docs.get(top_doc, {})
        short_sum = doc_data.get("summary10", "")
        long_sum = doc_data.get("summary100", "")

        # Filter summaries by keyword
        keyword = user_q.lower()
        filtered_short = " ".join([s for s in short_sum.split(". ") if keyword in s.lower()])
        filtered_long = " ".join([s for s in long_sum.split(". ") if keyword in s.lower()])

        st.markdown("**Short Summary (10%) filtered by keyword:**")
        st.text_area("Summary 10%", value=filtered_short or short_sum, height=200)

        st.markdown("**Detailed Summary (100%) filtered by keyword:**")
        st.text_area("Summary 100%", value=filtered_long or long_sum, height=300)

st.divider()

if st.session_state.docs:
    st.subheader("📚 Uploaded Documents")
    tabs = st.tabs(list(st.session_state.docs.keys()))
    for tab, name in zip(tabs, st.session_state.docs.keys()):
        with tab:
            st.text_area("Full Extracted Text", value=st.session_state.docs[name]["text"], height=300)
            st.download_button("⬇️ Download Extracted Text", data=st.session_state.docs[name]["text"], file_name=f"{name}_extracted.txt", mime="text/plain")
>>>>>>> 4f5bf9a8a33864fcb24b4b71fdf03d72ca47e619
